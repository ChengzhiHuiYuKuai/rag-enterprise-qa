"""多格式文档加载器"""

from pathlib import Path
from loguru import logger

from langchain_core.documents import Document


class DocumentLoader:
    """统一的文档加载接口，支持 PDF/Word/Markdown/Excel"""

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".md", ".xlsx", ".xls"}

    @classmethod
    def is_supported(cls, filename: str) -> bool:
        return Path(filename).suffix.lower() in cls.SUPPORTED_EXTENSIONS

    @classmethod
    def load(cls, file_path: str) -> list[Document]:
        """根据文件扩展名自动选择加载器"""
        path = Path(file_path)
        ext = path.suffix.lower()

        if not cls.is_supported(path.name):
            raise ValueError(f"不支持的文件格式: {ext}")

        loader_map = {
            ".pdf": cls._load_pdf,
            ".docx": cls._load_docx,
            ".md": cls._load_markdown,
            ".xlsx": cls._load_excel,
            ".xls": cls._load_excel,
        }

        loader = loader_map[ext]
        docs = loader(path)
        logger.info(f"加载文件 {path.name}，得到 {len(docs)} 个文档片段")
        return docs

    @staticmethod
    def _load_pdf(path: Path) -> list[Document]:
        """加载 PDF，使用 PyPDF 提取文本"""
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        docs = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text and text.strip():
                docs.append(
                    Document(
                        page_content=text.strip(),
                        metadata={
                            "source": path.name,
                            "page": i + 1,
                            "total_pages": len(reader.pages),
                            "file_type": "pdf",
                        },
                    )
                )
        return docs

    @staticmethod
    def _load_docx(path: Path) -> list[Document]:
        """加载 Word 文档"""
        from docx import Document as DocxDocument

        doc = DocxDocument(str(path))
        docs = []

        # 按段落提取，保留标题层级信息
        current_section = ""
        paragraphs = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # 检测标题
            if para.style.name.startswith("Heading"):
                # 先保存之前的段落
                if paragraphs:
                    content = f"{current_section}\n\n" if current_section else ""
                    content += "\n".join(paragraphs)
                    docs.append(
                        Document(
                            page_content=content,
                            metadata={
                                "source": path.name,
                                "section": current_section,
                                "file_type": "docx",
                            },
                        )
                    )
                    paragraphs = []
                current_section = text
            else:
                paragraphs.append(text)

        # 保存最后一段
        if paragraphs:
            content = f"{current_section}\n\n" if current_section else ""
            content += "\n".join(paragraphs)
            docs.append(
                Document(
                    page_content=content,
                    metadata={
                        "source": path.name,
                        "section": current_section,
                        "file_type": "docx",
                    },
                )
            )

        # 如果文档没有标题结构，整体作为一个文档
        if not docs and doc.paragraphs:
            full_text = "\n".join(
                p.text.strip() for p in doc.paragraphs if p.text.strip()
            )
            if full_text:
                docs.append(
                    Document(
                        page_content=full_text,
                        metadata={"source": path.name, "file_type": "docx"},
                    )
                )

        return docs

    @staticmethod
    def _load_markdown(path: Path) -> list[Document]:
        """加载 Markdown 文件，按标题分段"""
        text = path.read_text(encoding="utf-8")

        # 按一级和二级标题分段
        sections = []
        current_title = ""
        current_content = []

        for line in text.split("\n"):
            if line.startswith("# ") or line.startswith("## "):
                if current_content:
                    sections.append(
                        Document(
                            page_content="\n".join(current_content).strip(),
                            metadata={
                                "source": path.name,
                                "section": current_title,
                                "file_type": "markdown",
                            },
                        )
                    )
                current_title = line.lstrip("# ").strip()
                current_content = [line]
            else:
                current_content.append(line)

        # 最后一段
        if current_content:
            sections.append(
                Document(
                    page_content="\n".join(current_content).strip(),
                    metadata={
                        "source": path.name,
                        "section": current_title,
                        "file_type": "markdown",
                    },
                )
            )

        # 如果没有标题结构，整体作为一个文档
        if not sections and text.strip():
            sections.append(
                Document(
                    page_content=text.strip(),
                    metadata={"source": path.name, "file_type": "markdown"},
                )
            )

        return sections

    @staticmethod
    def _load_excel(path: Path) -> list[Document]:
        """加载 Excel 文件，每个 sheet 转为文本表格"""
        from openpyxl import load_workbook

        wb = load_workbook(str(path), read_only=True, data_only=True)
        docs = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                # 过滤全空行
                if any(cell is not None for cell in row):
                    cells = [str(cell) if cell is not None else "" for cell in row]
                    rows.append(" | ".join(cells))

            if rows:
                content = f"[Sheet: {sheet_name}]\n" + "\n".join(rows)
                docs.append(
                    Document(
                        page_content=content,
                        metadata={
                            "source": path.name,
                            "sheet": sheet_name,
                            "file_type": "excel",
                        },
                    )
                )

        wb.close()
        return docs
